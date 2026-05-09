import streamlit as st
from audio_recorder_streamlit import audio_recorder
from deep_translator import GoogleTranslator
from gtts import gTTS
import speech_recognition as sr
from PIL import Image
import pytesseract
import cv2
import numpy as np
from io import BytesIO
import fitz
from docx import Document
import tempfile
import os

# ----------------------------
# CONFIG
# ----------------------------
st.set_page_config(
    page_title="PragyanAI Studio",
    layout="wide"
)

langs_dict = GoogleTranslator().get_supported_languages(
    as_dict=True
)


# ----------------------------
# HELPERS
# ----------------------------
def translate_text(text, target_code):

    if not text.strip():
        return ""

    translated = GoogleTranslator(
        source="auto",
        target=target_code
    ).translate(text)

    return translated


def create_audio(text, lang):

    tts = gTTS(
        text=text,
        lang=lang
    )

    audio_fp = BytesIO()

    tts.write_to_fp(audio_fp)

    audio_fp.seek(0)

    return audio_fp


# DOCX translation
def translate_docx(uploaded_doc, target_code):

    original_doc = Document(uploaded_doc)

    translated_doc = Document()

    # Paragraphs
    for para in original_doc.paragraphs:

        translated_text = translate_text(
            para.text,
            target_code
        )

        translated_doc.add_paragraph(
            translated_text
        )

    # Tables
    for table in original_doc.tables:

        new_table = translated_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )

        for i, row in enumerate(table.rows):

            for j, cell in enumerate(row.cells):

                translated_cell = translate_text(
                    cell.text,
                    target_code
                )

                new_table.cell(i, j).text = translated_cell

    output = BytesIO()

    translated_doc.save(output)

    output.seek(0)

    return output


# PDF extraction
def extract_text_from_pdf(pdf_bytes):

    text = ""

    with fitz.open(
        stream=pdf_bytes,
        filetype="pdf"
    ) as doc:

        for page in doc:
            text += page.get_text()

    return text


# ----------------------------
# MAIN
# ----------------------------
def main():

    st.title(
        "🌐 PragyanAI Multi-Modal Studio"
    )

    target_lang = st.sidebar.selectbox(
        "Target Language",
        list(langs_dict.keys())
    )

    target_code = langs_dict[target_lang]

    tabs = st.tabs([
        "📄 DOCX",
        "📸 Image/PDF",
        "🎤 Audio",
        "📝 Text"
    ])


    # ---------------- DOCX
    with tabs[0]:

        doc_file = st.file_uploader(
            "Upload DOCX",
            type=["docx"]
        )

        if doc_file:

            if st.button(
                "Translate DOCX"
            ):

                translated_doc = translate_docx(
                    doc_file,
                    target_code
                )

                st.success(
                    "DOCX translated"
                )

                st.download_button(
                    "Download DOCX",
                    translated_doc,
                    file_name="translated.docx"
                )


    # ---------------- IMAGE/PDF
    with tabs[1]:

        uploaded_file = st.file_uploader(
            "Upload Image or PDF",
            type=[
                "png",
                "jpg",
                "jpeg",
                "pdf"
            ]
        )

        if uploaded_file:

            if uploaded_file.type == "application/pdf":

                if st.button(
                    "Translate PDF"
                ):

                    extracted = extract_text_from_pdf(
                        uploaded_file.read()
                    )

                    translated = translate_text(
                        extracted,
                        target_code
                    )

                    doc = Document()

                    doc.add_paragraph(
                        translated
                    )

                    output = BytesIO()

                    doc.save(output)

                    output.seek(0)

                    st.download_button(
                        "Download DOCX",
                        output,
                        file_name="translated_pdf.docx"
                    )

            else:

                img = Image.open(
                    uploaded_file
                )

                st.image(
                    img,
                    width=300
                )

                if st.button(
                    "Translate Image"
                ):

                    extracted = pytesseract.image_to_string(
                        img
                    )

                    translated = translate_text(
                        extracted,
                        target_code
                    )

                    # Create docx
                    doc = Document()

                    # image stays in doc
                    uploaded_file.seek(0)
                    doc.add_picture(
                        uploaded_file
                    )

                    doc.add_paragraph(
                        translated
                    )

                    output = BytesIO()

                    doc.save(output)

                    output.seek(0)

                    st.download_button(
                        "Download DOCX",
                        output,
                        file_name="translated_image.docx"
                    )


    # ---------------- AUDIO
    with tabs[2]:

        choice = st.radio(
            "Input type",
            [
                "Record",
                "Upload"
            ]
        )

        audio_bytes = None

        if choice == "Record":

            audio_bytes = audio_recorder(
                text="Record"
            )

        else:

            uploaded_audio = st.file_uploader(
                "Upload audio",
                type=[
                    "wav",
                    "flac"
                ]
            )

            if uploaded_audio:
                audio_bytes = uploaded_audio.read()

        if audio_bytes:

            st.audio(
                audio_bytes
            )

            if st.button(
                "Translate Audio"
            ):

                recognizer = sr.Recognizer()

                with sr.AudioFile(
                    BytesIO(audio_bytes)
                ) as source:

                    audio_data = recognizer.record(
                        source
                    )

                text = recognizer.recognize_google(
                    audio_data
                )

                translated = translate_text(
                    text,
                    target_code
                )

                st.success(
                    translated
                )

                audio_output = create_audio(
                    translated,
                    target_code
                )

                st.audio(
                    audio_output
                )


    # ---------------- TEXT
    with tabs[3]:

        text = st.text_area(
            "Type text"
        )

        if st.button(
            "Translate Text"
        ):

            translated = translate_text(
                text,
                target_code
            )

            st.success(
                translated
            )

            # Save as DOCX
            doc = Document()

            doc.add_paragraph(
                translated
            )

            output = BytesIO()

            doc.save(output)

            output.seek(0)

            st.download_button(
                "Download DOCX",
                output,
                file_name="translated_text.docx"
            )


if __name__ == "__main__":
    main()
